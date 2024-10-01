from http import client
from docx import Document
import pandas as pd
import glob
import os
from openai import OpenAI
import time
import random
import math
import numpy as np

def title_doi(foldpath):
    excel_files = glob.glob(os.path.join(foldpath, '*.xls'))
    print("hello")
    # print(excel_files)
    data = []
    # 遍历找到的Excel文件
    for file in excel_files:
        # 使用pandas读取Excel文件
        df = pd.read_excel(file, usecols=[11,32,39,42], header = None, skiprows=1)
        data.append(df)
        # 打印文件名和数据框的前几行
 
    # 将列表中的所有数据框合并成一个大的数据框
    combined_data = pd.concat(data, ignore_index=True)

    # 将数据框转换为矩阵A
    datalist = combined_data.values
    return datalist

def tell_kimi(model_text,system_text,user_text,temperature_number):
    # "sk-cGCkiSNwp1TkxDo6xLXMtCCgUhgK6CyT1G5e8KaCruQjrdiD"
    completion = client.chat.completions.create(
        # 传入模型参数，moonshot-v1是训练用于理解自然语言和书面语言的
        # moonshot-v1-8k: 它是一个长度为 8k 的模型，适用于生成短文本
        # moonshot-v1-32k: 它是一个长度为 32k 的模型，适用于生成长文本
        # moonshot-v1-128k: 它是一个长度为 128k 的模型，适用于生成超长文本
        model = model_text,
        # 输入文本
        messages = [
        # 系统消息
        {"role": "system", "content": system_text},
        # 用户消息
        {"role": "user", "content": user_text}
        ],
        # 确定结果的随机性，较高的值（如 0.7）将使输出更加随机，而较低的值（如 0.3）将使其更加集中和确定性
        temperature = temperature_number,
        )
    # 返回结果
    return(completion.choices[0].message.content)






foldpath = '存储wos检索结果excel的文件夹'
datalist = title_doi(foldpath)
outputfilepath = './Doc1.docx'
doc = Document(outputfilepath)


client = OpenAI(
    api_key = "替换为自己的api",
    base_url = "https://api.moonshot.cn/v1",
)

##循环前准备
# 获取需要的循环数
nmax=len(datalist)
# 设置系统模型
model_text='moonshot-v1-8k'
# 设置系统温度
temperature_number=0.3
# 设置系统文本
system_text='你是Kimi，由Moonshot AI提供的人工智能助手，你更擅长中文和英文的对话。你会为用户提供学术性强，有帮助，准确的回答。Moonshot AI为专有名词，不可翻译成其他语言。'
new_column = np.random.rand(nmax, 1)
new_column = new_column.astype(object)

for n in range(265, nmax):
    # if n == 99:
    #     break
    num = n + 1
    title = datalist[n, 0]
    abstract = datalist[n, 3]
    if isinstance(abstract, float):
        if math.isnan(abstract):
            continue
    add = "帮我将下面这篇文献的题目和摘要翻译成中文\n"
    user_text = add +"这是题目" + title +"这是摘要"+ abstract
    text = "第 " + str(num) + " 篇文章\n"
    p1 = doc.add_paragraph(text + title + "\n\n" + abstract + "\n")
    print("开始处理第" + str(num) + "篇文献\n")

    # 防止出现报错
    try:
        # 没有报错时生成比对文本
        report_text = tell_kimi(model_text, system_text, user_text, temperature_number)
        new_column[n, 0] = report_text

        # 添加一个段落
        doc.add_paragraph(report_text + "\n\n")
        # 添加分页符
        doc.add_page_break()
        doc.save(outputfilepath)
        # # 生成一个1到10之间的随机整数，表示睡眠的秒数
        # seconds = random.randint(1, 5)
        # # 睡眠指定的秒数
        # time.sleep(seconds)
        # #print("hello")
    except:
        # 生成报错信息
        error_message = "处理第 " + str(num) + " 时文献，kimi发生了错误，跳出循环"
        # 打印报错信息
        print(error_message)
        # 储存保存信息
        with open('error_log.txt', 'a') as file:
            file.write(error_message + '\n')
        # 继续执行下一个循环迭代
        continue
 # 保存到.docx文件

extended_array = np.concatenate((datalist, new_column), axis=1)

# 将NumPy数组转换为pandas DataFrame
columns = ['Column1', 'Column2', 'Column3', 'Column4', 'Column5']

# 将NumPy数组转换为pandas DataFrame
df = pd.DataFrame(extended_array, columns=columns)

# 将DataFrame保存为CSV文件
df.to_csv('output.csv', encoding='utf-8-sig',index=False)




print("hello")




